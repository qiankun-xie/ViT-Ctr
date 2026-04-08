#!/usr/bin/env python
"""
Generate Word (.docx) manuscripts from LaTeX source.
Produces:
  - paper/manuscript_en.docx  (English)
  - paper/manuscript_cn.docx  (Chinese)

Strategy:
  1. Use pandoc to convert LaTeX -> English docx (preserves equations & figures)
  2. Use python-docx to build Chinese version with translated content

Usage: python scripts/generate_docx.py
"""
import os
import re
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
PAPER_DIR = ROOT / "paper"
FIGURES_DIR = ROOT / "figures"
TEX_FILE = PAPER_DIR / "manuscript.tex"
EN_DOCX = PAPER_DIR / "manuscript_en.docx"
CN_DOCX = PAPER_DIR / "manuscript_cn.docx"


# ── Step 1: English version via pandoc ──────────────────────────────
def generate_english():
    """Build English docx with python-docx (includes embedded figures)."""
    print("[EN] Generating English Word manuscript via python-docx ...")
    doc = Document()
    sections_data = parse_latex()
    _apply_style(doc)

    # Title
    p = doc.add_heading(sections_data["title"], level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    p = doc.add_paragraph(sections_data["authors"])
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("Department of Chemistry, University, City, Country")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for sec in sections_data["sections"]:
        doc.add_heading(sec["heading"], level=sec["level"])
        for block in sec["blocks"]:
            if block["type"] == "text":
                doc.add_paragraph(block["content"])
            elif block["type"] == "figure":
                _add_figure(doc, block["path"], block.get("caption", ""))

    # References placeholder
    doc.add_heading("References", level=1)
    doc.add_paragraph("See paper/references.bib for the complete bibliography.")

    doc.save(str(EN_DOCX))
    size = EN_DOCX.stat().st_size
    print(f"  Created {EN_DOCX.name} ({size:,} bytes)")


def generate_english_pydocx():
    """Fallback: build English docx with python-docx from parsed LaTeX."""
    doc = Document()
    sections_data = parse_latex()
    _apply_style(doc)
    doc.add_heading(sections_data["title"], level=0)
    doc.add_paragraph(sections_data["authors"])
    for sec in sections_data["sections"]:
        doc.add_heading(sec["heading"], level=sec["level"])
        for block in sec["blocks"]:
            if block["type"] == "text":
                doc.add_paragraph(block["content"])
            elif block["type"] == "figure":
                _add_figure(doc, block["path"], block.get("caption", ""))
    doc.save(str(EN_DOCX))
    print(f"  Created {EN_DOCX.name} ({EN_DOCX.stat().st_size:,} bytes)")


# ── Step 2: Chinese version via python-docx ─────────────────────────
# Section heading translations
HEADING_MAP = {
    "Introduction": "引言",
    "Computational Methods": "计算方法",
    "Results and Discussion": "结果与讨论",
    "Conclusions": "结论",
    "Acknowledgement": "致谢",
    "Supporting Information": "支持信息",
    # Subsections
    "RAFT Kinetic ODE Model": "RAFT动力学ODE模型",
    "Chain Transfer Fingerprint (ctFP) Encoding": "链转移指纹（ctFP）编码",
    "Dataset Generation": "数据集生成",
    "SimpViT Architecture": "SimpViT架构",
    "Training Protocol": "训练方案",
    "Bootstrap Uncertainty Quantification": "Bootstrap不确定性量化",
    "Literature Validation Protocol": "文献验证方案",
    "Model Performance on Synthetic Test Set": "合成测试集上的模型性能",
    "Literature Validation: ML vs Mayo Baseline": "文献验证：ML与Mayo基线对比",
    "Per-RAFT-Agent-Class Analysis": "按RAFT试剂类别分析",
    "Uncertainty Quantification Assessment": "不确定性量化评估",
    "Training Dynamics": "训练动态",
    "Limitations": "局限性",
    "Web Application": "Web应用",
}

# Body text translations (section-level)
BODY_CN = {
    "abstract": (
        "可逆加成-断裂链转移（RAFT）聚合是最通用的可控自由基聚合技术之一，"
        "但传统上表征一个RAFT试剂需要三套独立的实验方案来分别测定链转移常数"
        "（$C_{tr}$）、抑制期和阻滞因子。本文提出一种深度学习方法，可从单组"
        "标准动力学数据——数均分子量（$M_n$）和分散度（Đ）随单体转化率在不同"
        "CTA/单体比下的变化——同时提取这三个参数。实验数据被编码为双通道链转移"
        "指纹（ctFP），一个64×64×2的图像张量，由简化Vision Transformer"
        "（SimpViT，约877,000参数）处理。模型在约973,000个合成样本上训练，"
        "这些样本通过对RAFT动力学常微分方程进行数值积分生成，涵盖四类RAFT试剂"
        "（二硫酯、三硫代碳酸酯、黄原酸酯和二硫代氨基甲酸酯）。采用200个微调头"
        "的Bootstrap不确定性量化和F分布联合置信区间提供校准的预测不确定性。"
        "对77个已发表的$C_{tr}$值的验证结果为$R^2 = 0.97$，"
        "RMSE($\\log_{10}$) = 0.18，92.2%的预测在已发表值的2倍范围内，"
        "100%在10倍范围内，显著优于Mayo方程ODE拟合基线"
        "（$R^2 = 0.50$，RMSE = 0.72）。提供免费Web应用供社区使用。"
    ),
    "introduction": (
        "可逆加成-断裂链转移（RAFT）聚合已成为应用最广泛的可控自由基聚合（CRP）"
        "技术之一，具有广泛的单体兼容性、温和的反应条件和对多种官能团的耐受性。"
        "RAFT试剂的有效性由多个动力学参数决定，其中链转移常数"
        "$C_{tr} = k_{add}/k_p$是描述RAFT试剂反应活性最基本的参数。"
        "此外，抑制期——聚合达到稳态前的初始延迟——和阻滞因子——"
        "有无RAFT试剂时聚合速率之比——分别提供了关于预平衡和主平衡动力学的关键信息。\n\n"
        "传统上，这三个参数各需要专门的实验方案。链转移常数最常通过Mayo方法测定，"
        "需要在不同CTA浓度下进行一系列低转化率聚合，然后分析所得分子量分布。"
        "抑制期从转化率-时间曲线中提取，阻滞因子从RAFT聚合与常规聚合的速率比较中获得。"
        "对于单个RAFT试剂-单体对，完整表征至少需要三组独立实验。\n\n"
        "机器学习在高分子科学中的最新进展表明，深度学习模型可以直接从实验观测量中"
        "提取动力学参数。特别是ViT-RR方法表明，在模拟共聚数据上训练的Vision "
        "Transformer，以反应活性指纹编码，可以预测反应活性比，精度可与传统非线性"
        "拟合方法相当或更优。\n\n"
        "本工作将这一范式扩展到RAFT聚合，证明单个模型可以从标准动力学数据"
        "（$M_n$和Đ随转化率在多个CTA/单体比下的变化）同时提取$C_{tr}$、"
        "抑制期和阻滞因子。关键洞察是：编码$M_n$和Đ作为CTA比和转化率函数的"
        "双通道链转移指纹（ctFP）包含足够信息来同时解析这三个参数。\n\n"
        "我们的方法有三个主要贡献：\n"
        "1. 涵盖四类RAFT试剂的综合ODE模型，包括二硫酯的两阶段预平衡机制，"
        "用于生成约973,000个合成训练样本。\n"
        "2. 处理ctFP图像同时预测三个参数的SimpViT，配合Bootstrap不确定性量化。\n"
        "3. 对来自三个独立文献来源的77个已发表$C_{tr}$值的验证，"
        "$R^2 = 0.97$，显著优于Mayo基线（$R^2 = 0.50$）。\n\n"
        "提供免费Web应用，使研究人员无需编程即可获得带置信区间的预测。"
    ),
    "methods": (
        "【计算方法部分包含大量数学公式和技术细节，请参阅英文版manuscript_en.docx"
        "或LaTeX源文件paper/manuscript.tex获取完整内容。以下为各小节概要。】\n\n"
        "RAFT动力学ODE模型：采用矩方法对完整RAFT机理建模，跟踪活性自由基链（μ）、"
        "休眠RAFT封端链（ν）和终止死链（λ）的零阶、一阶和二阶矩。"
        "三硫代碳酸酯、黄原酸酯和二硫代氨基甲酸酯采用14变量单平衡模型；"
        "二硫酯采用16变量预平衡模型。ODE系统使用scipy的Radau方法积分。\n\n"
        "链转移指纹（ctFP）编码：动力学数据编码为64×64×2的双通道图像张量。"
        "通道0为归一化$M_n$，通道1为分散度（截断于4.0）。\n\n"
        "数据集生成：通过拉丁超立方采样在7维参数空间上生成训练数据，"
        "共约973,000个有效样本，按80/10/10划分为训练/验证/测试集。\n\n"
        "SimpViT架构：输入64×64×2，patch大小16×16（16个patch），"
        "隐藏维度64，2层Transformer编码器，4个注意力头，约877,000参数。\n\n"
        "训练方案：Adam优化器，学习率3×10⁻⁴，批大小256，"
        "训练142个epoch，最佳模型在第126个epoch。\n\n"
        "Bootstrap不确定性量化：冻结骨干网络，微调200个独立输出头，"
        "使用F分布计算95%联合置信区间。校准因子分别为[100.0, 53.7, 3.5]。\n\n"
        "文献验证方案：对77个已发表$C_{tr}$值进行验证，"
        "每个数据点生成50个ML预测的集成，取中位数作为最终估计。"
    ),
    "results": (
        "合成测试集性能：训练后的SimpViT模型在三个输出参数上均表现优异。"
        "等值线图显示$\\log_{10}(C_{tr})$、抑制期和阻滞因子的预测值"
        "紧密聚集在恒等线附近。\n\n"
        "文献验证——ML与Mayo基线对比：对77个已发表$C_{tr}$值的验证表明，"
        "SimpViT模型达到$R^2 = 0.968$，RMSE($\\log_{10}$) = 0.181，"
        "中位倍数误差1.10。92.2%的ML预测在已发表值的2倍范围内，100%在10倍范围内。"
        "相比之下，Mayo ODE拟合基线$R^2 = 0.502$，RMSE = 0.715，"
        "85.7%在2倍范围内，90.9%在10倍范围内。\n\n"
        "值得注意的是，此处的Mayo基线使用固定的平均动力学参数，"
        "这代表了一种理想化场景。实际上研究人员必须独立估计这些参数，"
        "引入额外不确定性。ML模型甚至优于这一理想化基线，"
        "凸显了学习表征的鲁棒性。\n\n"
        "按RAFT试剂类别分析：二硫酯由于其预平衡动力学表现出最显著的"
        "抑制和阻滞效应，模型对其捕获良好。对于三硫代碳酸酯、黄原酸酯和"
        "二硫代氨基甲酸酯，预测的阻滞因子约为1.0，与已知化学性质一致。\n\n"
        "不确定性量化评估：校准后，抑制期和阻滞因子的95%置信区间覆盖率"
        "达到标称水平（95.0%）。但$\\log_{10}(C_{tr})$覆盖率仅为69.2%，"
        "表明Bootstrap集成系统性低估了该输出的不确定性。\n\n"
        "局限性：(1) 模拟与现实的差距；(2) 阻滞预测范围主要适用于二硫酯；"
        "(3) 参数范围边界外的外推未经验证；(4) 未涵盖温度依赖性和溶剂效应。\n\n"
        "Web应用：开发了免费的Streamlit Web应用，用户可输入实验数据"
        "获得带95%置信区间的预测结果。"
    ),
    "conclusions": (
        "我们证明了一个简化的Vision Transformer，在约973,000个ODE模拟的"
        "RAFT聚合曲线（编码为链转移指纹）上训练后，可以从标准动力学数据中"
        "同时提取链转移常数（$C_{tr}$）、抑制期和阻滞因子。这代表了重要的"
        "实际进步：传统方法需要三套独立实验方案分别测定这些参数，"
        "而我们的方法只需要一组$M_n$和Đ随转化率在不同CTA/单体比下的数据。\n\n"
        "对77个已发表$C_{tr}$值（涵盖四类RAFT试剂）的验证结果为"
        "$R^2 = 0.97$，92.2%的预测在已发表值的2倍范围内，"
        "显著优于Mayo方程ODE拟合基线（$R^2 = 0.50$）。\n\n"
        "三参数同时提取对二硫酯RAFT试剂最具信息量，其抑制和阻滞效应"
        "在动力学上具有重要意义。对于其他RAFT试剂类别，"
        "模型正确预测了最小阻滞，与已建立的RAFT化学一致。\n\n"
        "未来方向：一个有趣的扩展是直接从分子结构预测$C_{tr}$。"
        "通过将RAFT试剂结构编码为SMILES字符串并使用分子指纹或图神经网络"
        "作为输入表示，可能无需任何聚合实验即可预测链转移活性。"
        "这种结构-活性模型（Route A）可作为RAFT试剂设计的快速筛选工具，"
        "补充本文提出的实验数据驱动方法（Route B）。"
    ),
}

# Figure captions in Chinese
FIGURE_CAPTIONS_CN = {
    "fig1_concept.png": "示意工作流程：实验RAFT动力学数据编码为双通道ctFP图像，经SimpViT处理，解码为三个动力学参数及Bootstrap置信区间。",
    "fig2_ctfp_example.png": "代表性二硫酯RAFT试剂的链转移指纹（ctFP）示例。通道0（左）：归一化$M_n/M_{n,theory}$。通道1（右）：分散度Đ（截断于4.0）。",
    "fig3_parity_composite.png": "合成测试集上三个预测参数的等值线图：(a) $\\log_{10}(C_{tr})$，(b) 抑制期，(c) 阻滞因子。",
    "parity_ml_vs_mayo.png": "文献验证：ML预测与Mayo ODE拟合基线对77个已发表$C_{tr}$值的比较。",
    "dithioester_log10_Ctr.png": "二硫酯类$\\log_{10}(C_{tr})$预测的等值线图。",
    "loss_curves.png": "142个epoch的训练和验证损失曲线。最佳模型（第126个epoch）已标注。",
}


def generate_chinese():
    """Build Chinese Word manuscript with python-docx."""
    print("[CN] Generating Chinese Word manuscript ...")
    doc = Document()
    _apply_style(doc)

    # Title
    title_cn = (
        "利用Vision Transformer从RAFT聚合动力学数据中"
        "同时提取链转移常数、抑制期和阻滞因子"
    )
    p = doc.add_heading(title_cn, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Authors
    p = doc.add_paragraph("作者一, 作者二*")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("化学系，大学，城市，国家")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    # Abstract
    doc.add_heading("摘要", level=1)
    doc.add_paragraph(BODY_CN["abstract"])

    # Introduction
    doc.add_heading("引言", level=1)
    doc.add_paragraph(BODY_CN["introduction"])

    # Concept figure
    _add_figure(doc, FIGURES_DIR / "fig1_concept.png",
                FIGURE_CAPTIONS_CN.get("fig1_concept.png", ""))
    _add_figure(doc, FIGURES_DIR / "fig2_ctfp_example.png",
                FIGURE_CAPTIONS_CN.get("fig2_ctfp_example.png", ""))

    # Methods
    doc.add_heading("计算方法", level=1)
    doc.add_paragraph(BODY_CN["methods"])

    # Results
    doc.add_heading("结果与讨论", level=1)
    doc.add_paragraph(BODY_CN["results"])

    # Figures in results
    _add_figure(doc, FIGURES_DIR / "fig3_parity_composite.png",
                FIGURE_CAPTIONS_CN.get("fig3_parity_composite.png", ""))
    _add_figure(doc, FIGURES_DIR / "validation" / "parity_ml_vs_mayo.png",
                FIGURE_CAPTIONS_CN.get("parity_ml_vs_mayo.png", ""))
    _add_figure(doc, FIGURES_DIR / "parity_by_class" / "dithioester_log10_Ctr.png",
                FIGURE_CAPTIONS_CN.get("dithioester_log10_Ctr.png", ""))
    _add_figure(doc, FIGURES_DIR / "loss_curves.png",
                FIGURE_CAPTIONS_CN.get("loss_curves.png", ""))

    # Conclusions
    doc.add_heading("结论", level=1)
    doc.add_paragraph(BODY_CN["conclusions"])

    # Acknowledgements
    doc.add_heading("致谢", level=1)
    doc.add_paragraph("作者感谢[机构]提供的计算资源。[其他致谢待补充。]")

    # References placeholder
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[参考文献列表请参阅英文版或LaTeX源文件中的references.bib。]")

    doc.save(str(CN_DOCX))
    size = CN_DOCX.stat().st_size
    print(f"  Created {CN_DOCX.name} ({size:,} bytes)")


# ── Helpers ─────────────────────────────────────────────────────────
def _apply_style(doc):
    """Set default font and spacing."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5


def _add_figure(doc, path, caption=""):
    """Add a figure with caption if the file exists."""
    path = Path(path)
    if not path.exists():
        doc.add_paragraph(f"[Figure not found: {path.name}]")
        return
    try:
        doc.add_picture(str(path), width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Could not embed figure {path.name}: {e}]")
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)
            run.font.italic = True


def parse_latex():
    """Minimal LaTeX parser for fallback English docx generation."""
    tex = TEX_FILE.read_text(encoding="utf-8")

    # Extract title
    m = re.search(r"\\title\{(.+?)\}", tex, re.DOTALL)
    title = _strip_latex(m.group(1)) if m else "Manuscript"

    # Extract authors
    authors = []
    for m in re.finditer(r"\\author\{(.+?)\}", tex):
        authors.append(_strip_latex(m.group(1)))
    authors_str = ", ".join(authors)

    # Extract sections
    sections = []
    # Split by \section and \subsection
    parts = re.split(r"\\((?:sub)?section)\{(.+?)\}", tex)
    i = 1
    while i < len(parts):
        cmd = parts[i]
        heading = _strip_latex(parts[i + 1])
        body = parts[i + 2] if i + 2 < len(parts) else ""
        level = 1 if cmd == "section" else 2
        blocks = _parse_blocks(body)
        sections.append({"heading": heading, "level": level, "blocks": blocks})
        i += 3

    return {"title": title, "authors": authors_str, "sections": sections}


def _parse_blocks(text):
    """Parse text into text and figure blocks."""
    blocks = []
    parts = re.split(r"(\\begin\{figure\}.*?\\end\{figure\})", text, flags=re.DOTALL)
    for part in parts:
        if "\\begin{figure}" in part:
            m = re.search(r"\\includegraphics.*?\{(.+?)\}", part)
            cap = re.search(r"\\caption\{(.+?)\}", part, re.DOTALL)
            if m:
                fig_path = m.group(1).replace("../", "")
                blocks.append({
                    "type": "figure",
                    "path": str(ROOT / fig_path),
                    "caption": _strip_latex(cap.group(1)) if cap else "",
                })
        else:
            cleaned = _strip_latex(part).strip()
            if cleaned:
                blocks.append({"type": "text", "content": cleaned})
    return blocks


def _strip_latex(text):
    """Remove common LaTeX commands, keep readable text."""
    text = re.sub(r"\\cite\{[^}]*\}", "", text)
    text = re.sub(r"\\ref\{[^}]*\}", "X", text)
    text = re.sub(r"\\label\{[^}]*\}", "", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\texttt\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\mathrm\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\begin\{[^}]*\}", "", text)
    text = re.sub(r"\\end\{[^}]*\}", "", text)
    text = re.sub(r"\\item\s*", "• ", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = re.sub(r"[{}$]", "", text)
    text = re.sub(r"~", " ", text)
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"%%.*", "", text)
    return text.strip()


# ── Main ────────────────────────────────────────────────────────────
def main():
    PAPER_DIR.mkdir(parents=True, exist_ok=True)
    generate_english()
    generate_chinese()
    print("\nDone. Output files:")
    for f in [EN_DOCX, CN_DOCX]:
        if f.exists():
            print(f"  {f.relative_to(ROOT)}  ({f.stat().st_size:,} bytes)")
        else:
            print(f"  {f.relative_to(ROOT)}  MISSING")


if __name__ == "__main__":
    main()
